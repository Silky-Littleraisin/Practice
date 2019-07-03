# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from docx.shared import Inches

# 既存のdocxファイルを読み込む
doc = Document("basic.docx")

doc.add_page_break()  # 改ページ
doc.add_heading('図表を表示する', level=0)

doc.add_heading('図の配置', level=1)

doc.add_picture('wordcloud3.png', width=Inches(3.5))
last_p = doc.paragraphs[-1]
last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph("図1 ワードクラウド3", style='Caption')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_picture('wordcloud4.png', width=Inches(2.5))
last_p = doc.paragraphs[-1]
last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph("図2 ワードクラウド4", style='Caption')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()  # 改ページ


doc.add_heading('表の配置', level=1)

p = doc.add_paragraph("表1 2019年7月のカレンダー", style='Caption')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

records = (
    ('30', '1', '2', '3', '4', '5', '6',),
    ('7', '8', '9', '10', '11', '12', '13',),
    ('14', '15', '16', '17', '18', '19', '20',),
    ('21', '22', '23', '24', '25', '26', '27',),
    ('28', '29', '30', '31', '1', '2', '3',)
)

table = doc.add_table(rows=1, cols=7)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table.rows[0].cells
for i, w in enumerate(['日', '月', '火', '水', '木', '金', '土']):
    p = hdr_cells[i].add_paragraph(w)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if i == 0:
        p.runs[0].font.color.rgb = RGBColor(204, 0, 0)

for cols in records:
    row_cells = table.add_row().cells
    for i in range(0, 7):
        p = row_cells[i].add_paragraph(cols[i])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i == 0:
            p.runs[0].font.color.rgb = RGBColor(204, 0, 0)

doc.save('figure_table.docx')
