# -*- coding: utf-8 -*-
import docx

doc = docx.Document('basic.docx')

fullText = []
for para in doc.paragraphs:
    fullText.append(para.text)

print('\n'.join(fullText))

for para in doc.paragraphs:
    print(para.style.name, para.text, sep="\t")
