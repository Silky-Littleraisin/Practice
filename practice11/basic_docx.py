# -*- coding: utf-8 -*-
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import RGBColor
import datetime

now = datetime.datetime.today() # 現在の日にちを取得

doc = docx.Document()  # ドキュメントオブジェクトの作成

h = doc.add_heading('デジタルドキュメント論', level=0)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph(now.strftime("%Y年%m月%d日"))
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_heading('デジタルドキュメントとは', level=1)

p = doc.add_paragraph("デジタルドキュメントとは、新しい概念なので定まった定義はないが、デジタルドキュメント上で配信・流通・利用・保存されるドキュメントのことをいう。")

doc.add_heading('理解目標', level=2)
# 順序なしリスト
p = doc.add_paragraph('デジタルドキュメントの構成要素を理解する', style='List Bullet')
p.runs[0].bold = True
p = doc.add_paragraph('デジタルドキュメントを構成する技術要素を理解する', style='List Bullet')
p.runs[0].underline = True
p = doc.add_paragraph('デジタルドキュメントのコンピュータ処理の仕方を理解する', style='List Bullet')
p.runs[0].font.size = docx.shared.Pt(14)

doc.add_heading('講義スケジュール', level=2)
# 順序付きリスト
p = doc.add_paragraph('4/10 ガイダンス', style='List Number')
p.runs[0].font.color.rgb = RGBColor(204, 0, 0)
p = doc.add_paragraph('4/17 学術論文PDF解析', style='List Number')
p.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
p = doc.add_paragraph('4/24 小説の', style='List Number')
p.add_run('テキスト').font.color.rgb = RGBColor(0, 0, 230)
p.add_run('マイニング')

doc.save('basic.docx')